'''
Last modified: Tues Sep 20, 2022 @ 00:40
Last modified: Mon. Sep 26, 2022 @ 22:05

'''

from io import BytesIO
import streamlit as st
import time
from docx import Document
import pandas as pd
import xlsxwriter


def get_different_teams_to_list(uploadedFiles:list) -> dict:
    '''把桃園隊,執勤官, 其他國境隊各別歸到上述list'''
    taoyuan_teams_docx_list = []
    management_teams_excel_list = []
    other_teams_docx_list = []

    try:
        for file in uploadedFiles:
            if file.name.split(".")[-1] in ["docx","xlsx","xls"]:

                if file.name.endswith(".docx") and "桃園機場" in file.name:
                    taoyuan_teams_docx_list.append(file)
                elif file.name.endswith(".docx") and "特殊勤務" in file.name:
                    taoyuan_teams_docx_list.append(file)        
                elif "值勤官" in file.name:
                    management_teams_excel_list.append(file)
                elif "輪值表" in file.name:
                    if file.name.endswith(".docx") and file not in taoyuan_teams_docx_list :
                        other_teams_docx_list.append(file)
    except:
        raise ValueError("檔案格式有誤")
    
    mega_team = {
    "management_teams":management_teams_excel_list,
    "taoyuan_teams":taoyuan_teams_docx_list,
    "other_teams":other_teams_docx_list
    }
    
    return mega_team

def replace_with_names_xlsx(list_number:str, list_officer:str) -> list:
    '''Replace all indices with its corresponding officer names'''
    for i in range(len(list_number)):
        if list_number[i] == 0:
            list_number[i] = list_officer[0]
        elif list_number[i] == 1:
            list_number[i] = list_officer[1]
        elif list_number[i] == 2:
            list_number[i] = list_officer[2]
        elif list_number[i] == 3:
            list_number[i] = list_officer[3]
        elif list_number[i] == 4:
            list_number[i] = list_officer[4]
        elif list_number[i] == 5:
            list_number[i] = list_officer[5]
        elif list_number[i] == 6:
            list_number[i] = list_officer[6]
        elif list_number[i] == 7:
            list_number[i] = list_officer[7]

    return list_number

def replace_with_names_xls(list_number:str, list_officer:str) -> list:
    '''Replace all indices with its corresponding officer names for xls files'''
    for i in range(len(list_number)):
        if list_number[i] == 0:
            list_number[i] = list_officer[0]
        elif list_number[i] == 1:
            list_number[i] = list_officer[1]
        elif list_number[i] == 2:
            list_number[i] = list_officer[2]
        elif list_number[i] == 3:
            list_number[i] = list_officer[3]
        elif list_number[i] == 4:
            list_number[i] = list_officer[4]
        elif list_number[i] == 5:
            list_number[i] = list_officer[5]
        elif list_number[i] == 6:
            list_number[i] = list_officer[6]

    return list_number

def getDataFromExcel(teams: list) -> pd.DataFrame:
    for excelfile in teams:
        if excelfile.name.endswith(".xls"):
            df_xls = pd.read_excel(excelfile, header = None)
        elif excelfile.name.endswith(".xlsx"):
            df_xlsx = pd.read_excel(excelfile, header = None)

    # '''xls files'''
    header_xls = df_xls.iloc[3, 3:-1]
    df_xls = df_xls[:-1]
    df_xls1 = df_xls.iloc[4:, 3:-1].rename(columns = header_xls).reset_index(inplace=False, drop=True)

    ON_DUTY = '△'

    general_officer_list = []
    GeneralOfficerNames_list = df_xls1.columns.tolist()

    for row in range(len(df_xls1)):
        dalist = (df_xls1.iloc[row] == ON_DUTY).tolist()
        for idx in range(len(df_xls1.columns)):
            if dalist[idx]:
                general_officer_list.append(idx)

    general_officer_list = replace_with_names_xls(general_officer_list,GeneralOfficerNames_list)


    # '''xlsx files'''
    header_xlsx = df_xlsx.iloc[3,2:]
    df_xlsx = df_xlsx.iloc[4:-1,2:]
    df_xlsx = df_xlsx.rename(columns = header_xlsx).reset_index(inplace=False, drop=True)

    dep_officer_list = []
    DepOfficerNames_list = df_xlsx.columns.tolist()

    for row in range(len(df_xlsx)):
        dalist = (df_xlsx.iloc[row] == ON_DUTY).tolist()
        for idx in range(len(df_xlsx.columns)):
            if dalist[idx]:
                dep_officer_list.append(idx)

    dep_officer_list = replace_with_names_xlsx(dep_officer_list,DepOfficerNames_list)


    df_general_officer = pd.DataFrame(general_officer_list)
    df_dep_officer = pd.DataFrame(dep_officer_list)
    df_combined = pd.concat([df_general_officer,df_dep_officer], axis=1)

    df_combined.columns = ["總值勤官", "值勤官"]
    # add this to row[0]
    df_0 = pd.DataFrame("姓名",columns = ["總值勤官", "值勤官"], index=[0])
    df_combined = pd.concat([df_0,df_combined], axis=0).reset_index(drop=True)

    return df_combined

def get_DataFrame(team:str) -> pd.DataFrame:
    '''Read WordDoc table content to pd.DataFrame'''
    document = Document(team)
    table = document.tables[0]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)
    df.iloc[:,0] = df.iloc[:,0].str.strip().str.replace(" ","")

    # 隊名prefix
    prefix = team.name.split("月")[1][:5] + "_"
    if prefix[-2] == '一':
        prefix = "1"+prefix
    elif prefix[-2] == "二":
        prefix = "2" + prefix
    elif prefix[-2] == "三":
        prefix = "3" + prefix
    elif prefix[-2] == "四":
        prefix = "4" + prefix
    elif prefix[-2] == "五":
        prefix = "5" + prefix

    # 處理header
    for i in range(len(df)):
        if df.iloc[i,0] == "01":
            FirstDayRow = i
        elif df.iloc[i,0] == "1":
            FirstDayRow = i

    NameIndexRow = FirstDayRow -1
    NameIndexCol = 2

    titles = df.iloc[NameIndexRow,NameIndexCol:].copy()
    for i in range(NameIndexCol,len(titles)+2):
        titles[i] = (titles[i].split("\n")[0]).replace(" ","")

    header = df.iloc[NameIndexRow,NameIndexCol:].copy()
    for i in range(NameIndexCol,len(header)+2):
        header[i] = prefix + header[i].strip()[-3:]

    df_reorganized = df.iloc[FirstDayRow-1:-1,NameIndexCol:].reset_index(inplace=False, drop=True).rename(columns = header)
    if "/" in df_reorganized.iloc[-1,:].values:
        df_reorganized = df_reorganized.drop(df_reorganized.index[-1])
    df_reorganized.iloc[0,:] = titles
    return df_reorganized

def get_DataFrame_SpecialForce(team:str) -> pd.DataFrame:
    '''跟get_DataFrame一樣,但是特殊勤務隊的一號分隊長名稱怪異'''
    document = Document(team)
    table = document.tables[0]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)
    df.iloc[:,0] = df.iloc[:,0].str.strip().str.replace(" ","")

    # 隊名prefix
    prefix = "0"+team.name.split("月")[1][:5] + "_"

    # 處理header
    FirstDayRow = df.iloc[:,0].eq("01").idxmax()  
    NameIndexRow = FirstDayRow -1
    NameIndexCol = 2

    titles = df.iloc[NameIndexRow,NameIndexCol:].copy()
    for i in range(NameIndexCol,len(titles)+2):
        titles[i] = (titles[i].split("\n")[0]).replace(" ","")

    # "分隊長: 一洪金城代理 "
    header = df.iloc[NameIndexRow,NameIndexCol:].copy()
    header[2] = prefix + header[2][-3:] 
    header[3] = prefix + header[3][-2:]
    header[4] = prefix + header[4][1:4]
    header[5] = prefix + header[5][-3:]
    header[6] = prefix + header[6][-3:]

    ####
    df_reorganized = df.iloc[FirstDayRow-1:-1,NameIndexCol:].reset_index(inplace=False, drop=True).rename(columns = header)
    if "/" in df_reorganized.iloc[-1,:].values:
        df_reorganized = df_reorganized.drop(df_reorganized.index[-1])
    df_reorganized.iloc[0,:] = titles
    return df_reorganized

def get_DataFrame_KaoHsiungAirport(team:str) -> pd.DataFrame:
    '''跟get_DataFrame一樣,但是高雄機場隊有多一層隱藏column'''
    document = Document(team)
    table = document.tables[0]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)
    df.iloc[:,0] = df.iloc[:,0].str.strip().str.replace(" ","")

    # 處理多餘的column
    df = df.drop(df.columns[-1],axis=1)

    # 隊名prefix
    prefix = team.name.split("月")[1][:5] + "_"

    # 處理header
    FirstDayRow = df.iloc[:,0].eq("01").idxmax()  
    NameIndexRow = FirstDayRow -1
    NameIndexCol = 2

    titles = df.iloc[NameIndexRow,NameIndexCol:].copy()
    for i in range(NameIndexCol,len(titles)+2):
        titles[i] = (titles[i].split("\n")[0]).replace(" ","")

    header = df.iloc[NameIndexRow,NameIndexCol:].copy()
    for i in range(NameIndexCol,len(header)+2):
        header[i] = prefix + header[i].strip()[-3:]

    df_reorganized = df.iloc[FirstDayRow-1:-1,NameIndexCol:].reset_index(inplace=False, drop=True).rename(columns = header)
    if "/" in df_reorganized.iloc[-1,:].values:
        df_reorganized = df_reorganized.drop(df_reorganized.index[-1])
    df_reorganized.iloc[0,:] = titles

    return df_reorganized

def time_formatting(df:pd.DataFrame) -> pd.DataFrame:
    '''把時間字串整理整理'''
    for column in df.columns:
        df[column] = df[column].str.strip()

    df = df.replace("至","-", regex=True).replace("時","", regex=True).replace("前日","",regex=True).replace("▲","",regex=True)
    df = df.replace("*","").replace("△","",regex=True).replace("\(公\)","",regex=True)
    df = df.replace("■","",regex=True).replace("\(代理\)","",regex=True).replace(":","",regex=True)
    df = df.replace("前", "",regex=True).replace("9-21","09-21").replace("前20-8","20-8")

    return df

def get_day_of_week(excelpath:str) -> list:
    df = pd.read_excel(excelpath[0])
    df_day_of_week = df.iloc[2:-1,1].reset_index(drop=True)
    dayOfweek = df_day_of_week.tolist()

    return dayOfweek


def get_taoTeams(tao_teams:list) -> pd.DataFrame:
    '''把桃機隊+特勤隊放到一個df, 並做time formatting'''
    TaoCombined_df = pd.DataFrame()

    for tao_team in tao_teams:
        
        if "特殊勤務" in tao_team.name:
            df_taoyuan = get_DataFrame_SpecialForce(tao_team)
            TaoCombined_df = pd.concat([TaoCombined_df,df_taoyuan], axis=1) 

        else:
            df_taoyuan = get_DataFrame(tao_team)
            TaoCombined_df = pd.concat([TaoCombined_df,df_taoyuan], axis=1)

    TaoCombined_df = TaoCombined_df.replace("○","輪休").replace("","輪休")
    TaoCombined_df = time_formatting(TaoCombined_df)
    
    return TaoCombined_df

def get_otherTeams(other_teams:list) -> pd.DataFrame:
    '''把其他港隊機場隊放到一個df, 並做time formatting'''
    OtherCombined_df = pd.DataFrame()

    for other_team in other_teams:
        if "高雄機場" in other_team.name:
            df_other_team = get_DataFrame_KaoHsiungAirport(other_team)
            OtherCombined_df = pd.concat([OtherCombined_df,df_other_team], axis=1)
        else:
            df_other_team = get_DataFrame(other_team)
            OtherCombined_df = pd.concat([OtherCombined_df,df_other_team], axis=1)

    OtherCombined_df = OtherCombined_df.replace("○","輪休").replace("","輪休")
    OtherCombined_df = time_formatting(OtherCombined_df)
        
    return OtherCombined_df

def build_tao_spes_teams(tao_combined_df:pd.DataFrame):
    '''把tao1~tao5, specialForce 從tao_combined_df分開'''
    # 目前流程從把 各別list -> 統一 dict -> 統一df (time formatting)完後 ->  個別df 

    tao1_index = []
    tao2_index = []
    tao3_index = []
    tao4_index = []
    tao5_index = []
    specialForce_index = []
    
    for column_idx in range(len(tao_combined_df.columns.values)):

        if "1" in tao_combined_df.columns.values[column_idx]:
            tao1_index.append(column_idx)
            tao1_df = tao_combined_df.iloc[:, tao1_index[0]:tao1_index[-1]+1]

        elif "2" in tao_combined_df.columns.values[column_idx]:
            tao2_index.append(column_idx)
            tao2_df = tao_combined_df.iloc[:, tao2_index[0]:tao2_index[-1]+1]

        elif "3" in tao_combined_df.columns.values[column_idx]:
            tao3_index.append(column_idx)
            tao3_df = tao_combined_df.iloc[:, tao3_index[0]:tao3_index[-1]+1]

        elif "4" in tao_combined_df.columns.values[column_idx]:
            tao4_index.append(column_idx)
            tao4_df = tao_combined_df.iloc[:, tao4_index[0]:tao4_index[-1]+1]

        elif "5" in tao_combined_df.columns.values[column_idx]:
            tao5_index.append(column_idx)
            tao5_df = tao_combined_df.iloc[:, tao5_index[0]:tao5_index[-1]+1]

        elif "0" in tao_combined_df.columns.values[column_idx]:
            specialForce_index.append(column_idx)
            specialForce_df = tao_combined_df.iloc[:, specialForce_index[0]:specialForce_index[-1]+1]

    return tao1_df,tao2_df,tao3_df,tao4_df,tao5_df,specialForce_df

def build_other_teams(combined_teams: pd.DataFrame) -> pd.DataFrame:
    '''把所有其他港務機場門境隊 從other_combined_df分開'''
    keelung_index = []
    songshang_index = []
    taichung_index = []
    kaohsiungAirport_index = []
    kaohsiungPort_index = []
    jingmen_index = []

    for column_idx in range(len(combined_teams.columns.values)):

        if "基隆" in combined_teams.columns.values[column_idx]:
            keelung_index.append(column_idx)
            keelung_df = combined_teams.iloc[:, keelung_index[0]:keelung_index[-1]+1]

        elif "松山" in combined_teams.columns.values[column_idx]:
            songshang_index.append(column_idx)
            songshang_df = combined_teams.iloc[:, songshang_index[0]:songshang_index[-1]+1]

        elif "臺中" in combined_teams.columns.values[column_idx]:
            taichung_index.append(column_idx)
            taichung_df = combined_teams.iloc[:, taichung_index[0]:taichung_index[-1]+1]

        elif "高雄機場" in combined_teams.columns.values[column_idx]:
            kaohsiungAirport_index.append(column_idx)
            kaohsiungAirport_df = combined_teams.iloc[:, kaohsiungAirport_index[0]:kaohsiungAirport_index[-1]+1]

        elif "高雄港隊" in combined_teams.columns.values[column_idx]:
            kaohsiungPort_index.append(column_idx)
            kaohsiungPort_df = combined_teams.iloc[:, kaohsiungPort_index[0]:kaohsiungPort_index[-1]+1]

        elif "金門" in combined_teams.columns.values[column_idx]:
            jingmen_index.append(column_idx)
            jingmen_df = combined_teams.iloc[:, jingmen_index[0]:jingmen_index[-1]+1]

    return keelung_df, songshang_df, taichung_df ,kaohsiungAirport_df, kaohsiungPort_df, jingmen_df



def get_my_excel_timetable(mega_team:dict, DAY:int,MONTH:str,YEAR:str):

    day_of_week = get_day_of_week(mega_team["management_teams"])
    df_officer = getDataFromExcel(mega_team["management_teams"])

    # '''TaoYuanCombined_df'''
    TaoYuanCombined_df = get_taoTeams(mega_team["taoyuan_teams"])

    tao1_df,tao2_df,tao3_df,tao4_df,tao5_df,specialForce_df = \
        build_tao_spes_teams(TaoYuanCombined_df)

    # '''OtherCombined_df'''
    OtherCombined_df = get_otherTeams(mega_team["other_teams"])

    keelung_df, songshang_df, taichung_df ,kaohsiungAirport_df, kaohsiungPort_df, \
        jingmen_df = build_other_teams(OtherCombined_df)

    excel_file_name = f'幹部出勤表_{MONTH}月{DAY}日.xlsx'

    output = BytesIO()
    with xlsxwriter.Workbook(output, {'in_memory': True}) as workbook:

        worksheet = workbook.add_worksheet("幹部出勤表")

        #'''CONSTANTS'''
        FIRST_ROW = 1
        FIRST_COL = 1
        LAST_COLUMNS = 8 # FIRST_COL+7
        LAST_ROWS = 37 # FIRST_ROW+36

        # By Cell
        first_row =  2
        first_col = "B"
        last_row = 38
        last_col = "I"

        # Border constants
        THICKEST = 5
        THICK = 2
        # Some useful row -> WHOLE ROW
        TitleWholeRow = first_col+str(first_row)+":"+last_col+str(first_row) #"B2:I2"
        CustomWholeTeamRow = "B3:I3"
        TaoYuanWholeRow = first_col+str(first_row+5)+":"+last_col+str(first_row+5)
        AirportPortWholeRow = first_col+str(first_row+22)+":"+last_col+str(first_row+22)
        NoteWholeRow = first_col+str(first_row+36)+":"+last_col+str(first_row+36)

        # Some useful row indices
        title_row = first_row
        custome_row = first_row + 1
        taoyuan_row = first_row + 5
        airportport_row = first_row+22
        note_row = last_row

        #'''=================================================='''
        #'''Format Defining'''
        # Font_color: black, font_size:14, bold:True
        subHeader_bold = workbook.add_format({'bold':True, "font_size":14,
                                            "align":"center","valign":"vcenter"})
        # text_format, align:center
        text_format_center_officers = workbook.add_format({"font_size":16, "align":"center",
                                                "valign":"vcenter"})
        # Add an Excel date format.
        date_format = workbook.add_format({'num_format':'mmmm d yyyy'})

        # row headers
        TitleRow_format = workbook.add_format({"bold":True, 'font_color':"white",'font_size':20,
                                            "align":"center","valign":"vcenter",
                                            "bg_color":"#1155CC","border":5})

        CustomTeamRow_format = workbook.add_format({"bold":True, 'font_color':"white",'font_size':16,
                                            "align":"center","valign":"vcenter",
                                            "bg_color":"#CC0000","border":5})

        TaoYuanRow_AirportPortRow_format = workbook.add_format({"bold":True, 'font_color':"white",'font_size':14,
                                            "align":"center","valign":"vcenter",
                                            "bg_color":"#CC0000","border":5})

        last_row_format = workbook.add_format({
                                                "bold":True,"font_color":"#C00000","font_size":14,
                                                "border":THICKEST,"top":THICK
                                            })

        grey_bg_color = workbook.add_format({"bg_color":"#D9D9D9"})

        # Border Thickness
        thickest_border = workbook.add_format({"border":THICKEST})
        thick_border = workbook.add_format({"border":THICK})



        #'''=================================================='''
        # Conditional Formatting
        def conditional_formatting():
            #'''Conditional formatting'''
            worksheet.conditional_format('B4:I6',{'type':'no_blanks','format':thick_border}) # thick_border in-between
            worksheet.conditional_format('B8:I23',{'type':'no_blanks','format':thick_border})
            worksheet.conditional_format('B8:I23',{'type':'blanks','format':thick_border})
            worksheet.conditional_format('B25:I37',{'type':'no_blanks','format':thick_border})
            worksheet.conditional_format('B25:I37',{'type':'blanks','format':thick_border})

            worksheet.conditional_format(FIRST_ROW,FIRST_COL-1,LAST_ROWS,FIRST_COL-1, {"type":"blanks","format":workbook.add_format({"right":THICKEST})})
            worksheet.conditional_format(FIRST_ROW,LAST_COLUMNS+1,LAST_ROWS,LAST_COLUMNS+1, {"type":"blanks","format":workbook.add_format({"left":THICKEST})})
            worksheet.conditional_format("I9:I37", {"type":"blanks","format":workbook.add_format({"left":THICK,"top":THICK,"bottom":THICK,"right":THICKEST})})

            # Individual cell Border formatting
            worksheet.conditional_format("B3:I3",{"type":"no_blanks","format":thickest_border})
            worksheet.conditional_format("B7:I7",{"type":"no_blanks","format":thickest_border})
            worksheet.conditional_format("B8:I8",{"type":"no_blanks","format":thickest_border})
            worksheet.conditional_format("B24:I24",{"type":"no_blanks","format":thickest_border})
            worksheet.conditional_format("B38:I38",{"type":"no_blanks","format":workbook.add_format({"top":THICK,"left":THICKEST,"right":THICKEST,"bottom":THICKEST})})

            #'''Row background_color Preset'''
            # worksheet.conditional_format('B11:I12',{'type':"text",'criteria':"not containing",
            #                                         "value":"-","format":grey_bg_color,
            #                                         "multi_range":"B11:I12 B15:I16 H19:I20 B25:I26 B30:I31 B34:I35"})
            # worksheet.conditional_format('B11:I12',{'type':"text",'criteria':"not containing",
            #                                     "value":"|","format":grey_bg_color,
            #                                     "multi_range":"B11:I12 B15:I16 H19:I20 B25:I26 B30:I31 B34:I35"})
            worksheet.conditional_format('B11:B12',{'type':"text",'criteria':"not containing",
                                                "value":"|","format":grey_bg_color,
                                                "multi_range":"B11:B12 D11:D12 F11:F12 H11:I12"})
            worksheet.conditional_format('B15:B16',{'type':"text",'criteria':"not containing",
                                                "value":"|","format":grey_bg_color,
                                                "multi_range":"B15:B16 D15:D16 F15:F16 H15:I16"})
            worksheet.conditional_format('H19:I20',{'type':"text",'criteria':"not containing",
                                                "value":"|","format":grey_bg_color})
            worksheet.conditional_format('B25:B27',{'type':"text",'criteria':"not containing",
                                                "value":"|","format":grey_bg_color,
                                                "multi_range":"B25:B27 D25:D27 F25:F27 H25:I27"})
            worksheet.conditional_format('B30:B31',{'type':"text",'criteria':"not containing",
                                                "value":"|","format":grey_bg_color,
                                                "multi_range":"B30:B31 D30:D31 F30:F31 H30:I31"})
            worksheet.conditional_format('B34:B35',{'type':"text",'criteria':"not containing",
                                                "value":"|","format":grey_bg_color,
                                                "multi_range":"B34:B35 D34:D35 F34:F35 H34:I35"})

        conditional_formatting()
        #'''=================================================='''
        #'''Cell widths'''
        def cell_width_adjusting():
            '''負責所有儲存格width adjusting'''

            worksheet.set_column(FIRST_COL,FIRST_COL, 12.33)
            worksheet.set_column(FIRST_COL+1,FIRST_COL+1, 4.83)
            worksheet.set_column(FIRST_COL+2,FIRST_COL+2, 8)
            worksheet.set_column(FIRST_COL+3,FIRST_COL+3, 4)
            worksheet.set_column(FIRST_COL+4,FIRST_COL+4, 8)
            worksheet.set_column(FIRST_COL+5,FIRST_COL+5, 9)
            worksheet.set_column(FIRST_COL+6,FIRST_COL+6, 13.17)
            worksheet.set_column(FIRST_COL+7,FIRST_COL+7, 12.17)

            #'''Cell heights'''
            for i in range(0,note_row):
                worksheet.set_row(i,23)
                if i == 10:
                    worksheet.set_row(i,23)
                    worksheet.set_row(i+1,23)
                elif i == 14:
                    worksheet.set_row(i,23, )
                    worksheet.set_row(i+1,23, )
                elif i == 18:
                    worksheet.set_row(i,23, )
                    worksheet.set_row(i+1,23, )
                elif i == 24:
                    worksheet.set_row(i,23, )
                    worksheet.set_row(i+1,23, )                        
                    worksheet.set_row(i+2,23, )
                elif i == 29:
                    worksheet.set_row(i,23, )
                    worksheet.set_row(i+1,23, )
                elif i == 33:
                    worksheet.set_row(i,23, )
                    worksheet.set_row(i+1,23, )      

            worksheet.set_row(10+1,23)
            worksheet.set_row(14+1,23, )
            worksheet.set_row(18+1,23, )
            worksheet.set_row(24+1,23, )
            worksheet.set_row(24+2,23, )
            worksheet.set_row(29+1,23, )
            worksheet.set_row(33+1,23, )

            worksheet.set_row(title_row-1,31)
            worksheet.set_row(custome_row-1,25)
            worksheet.set_row(custome_row,25)
            worksheet.set_row(custome_row+1,25)
            worksheet.set_row(taoyuan_row-1,23)
            worksheet.set_row(taoyuan_row,24)
            worksheet.set_row(airportport_row-1,23)
            worksheet.set_row(note_row-1,22)

        cell_width_adjusting()
        #'''=================================================='''
        def create_thick_cell_borders(border_thickness):
            '''Create surroudning thick borders'''
            for row_idx in range(FIRST_ROW, LAST_ROWS):
                for column_idx in range(FIRST_COL, LAST_COLUMNS+1):
                    worksheet.write(row_idx, column_idx, " ", border_thickness)

        create_thick_cell_borders(thick_border)
        #'''=================================================='''
        #'''固定不變的cells'''
        def table_construction():
            #'''建立固定不變的cells'''
            worksheet.merge_range(TitleWholeRow, f"中華民國 {YEAR} 年 {MONTH} 月 {DAY} 日 ( {day_of_week[DAY]} )", TitleRow_format)
            worksheet.merge_range(CustomWholeTeamRow, "國境事務大隊協管室值勤表", CustomTeamRow_format)
            #'''===================='''
            worksheet.write(first_col+str(custome_row+1), '總值勤官', workbook.add_format({'bold':True, "font_size":14,
                                                "align":"center","valign":"vcenter","top":5,"left":5,"right":2,"bottom":2}))
            worksheet.write(first_col+str(custome_row+2), '值勤官', workbook.add_format({'bold':True, "font_size":14,
                                                "align":"center","valign":"vcenter","top":2,"left":5,"right":2,"bottom":2}))
            worksheet.write(first_col+str(custome_row+3), '值勤員', workbook.add_format({'bold':True, "font_size":14,
                                                "align":"center","valign":"vcenter","top":2,"left":5,"right":2,"bottom":5}))
            worksheet.merge_range("C6:E6", "",workbook.add_format({'bold':False, "font_size":14,
                                                "align":"center","valign":"vcenter","top":2,"left":2,"right":2,"bottom":2})) #
            worksheet.merge_range("F6:G6", "監控人員",workbook.add_format({'bold':True, "font_size":14,
                                                "align":"center","valign":"vcenter","top":2,"left":2,"right":2,"bottom":5})) #
            worksheet.write(custome_row+2, LAST_COLUMNS-1, " ",workbook.add_format({'bold':False, "font_size":14,
                                                "align":"center","valign":"vcenter","top":2,"left":2,"right":2,"bottom":5})) 
            worksheet.write(custome_row+2, LAST_COLUMNS, " ",workbook.add_format({'bold':False, "font_size":14,
                                                "align":"center","valign":"vcenter","top":2,"left":2,"right":5,"bottom":2}))
            #'''===================='''
            worksheet.merge_range(TaoYuanWholeRow, '桃園機場幹部出勤狀況表', TaoYuanRow_AirportPortRow_format) 
            #'''===================='''
            worksheet.write(first_col+str(taoyuan_row+1), "隊別 / 職稱", subHeader_bold)
            worksheet.merge_range("C8:D8", "隊長", subHeader_bold) # merge
            worksheet.merge_range("E8:F8", "副隊長", subHeader_bold) # merge
            worksheet.merge_range("G8:H8", "分隊長", subHeader_bold) # merge
            worksheet.write(taoyuan_row,FIRST_COL+7, "備註", subHeader_bold) 

            taoyuan_airport_teams = [
                "桃機一隊", "桃機二隊","桃機四隊", "桃機五隊", "桃機三隊", "特殊勤務隊"
            ]

            for i in range(0,2*len(taoyuan_airport_teams),2):
                # i = 0,2,4,6,8,10
                j = int(i/2)
                # j = 0,1,2,3,4,5
                if taoyuan_airport_teams[j] not in ["桃機三隊", "特殊勤務隊"]:
                    worksheet.merge_range(first_col+f"{i+9}"+":"+first_col+f"{i+10}",taoyuan_airport_teams[j],subHeader_bold)
                elif taoyuan_airport_teams[j] in ["桃機三隊"]:
                    worksheet.merge_range(first_col+f"{i+9}"+":"+first_col+f"{i+12}",taoyuan_airport_teams[j],subHeader_bold)
                elif taoyuan_airport_teams[j] in ["特殊勤務隊"]:
                    worksheet.merge_range(first_col+f"{i+11}"+":"+first_col+f"{i+13}",taoyuan_airport_teams[j],subHeader_bold)
                else:
                    raise ValueError("Something went wrong")


            #'''===================='''
            worksheet.merge_range(AirportPortWholeRow, "外機港隊幹部出勤狀況表", TaoYuanRow_AirportPortRow_format) # merge, fontsize, white font, red background, border
            #'''===================='''
            other_port_teams = [
                "基隆港隊", "松山機場隊","臺中港隊", "高雄機場隊", "高雄港隊", "金門國境隊"
            ]

            worksheet.merge_range(first_col+f'{airportport_row+1}'+":"+first_col+f"{airportport_row+3}",other_port_teams[0],subHeader_bold)
            worksheet.merge_range(first_col+f'{airportport_row+4}'+":"+first_col+f"{airportport_row+5}",other_port_teams[1],subHeader_bold)
            worksheet.merge_range(first_col+f'{airportport_row+6}'+":"+first_col+f"{airportport_row+7}",other_port_teams[2],subHeader_bold)
            worksheet.merge_range(first_col+f'{airportport_row+8}'+":"+first_col+f"{airportport_row+9}",other_port_teams[3],subHeader_bold)
            worksheet.merge_range(first_col+f'{airportport_row+10}'+":"+first_col+f"{airportport_row+11}",other_port_teams[4],subHeader_bold)
            worksheet.merge_range(first_col+f'{airportport_row+12}'+":"+first_col+f"{airportport_row+13}",other_port_teams[5],subHeader_bold)

            #'''備註 Merge'''
            worksheet.merge_range(last_col+"9:"+last_col+"10"," ")
            worksheet.merge_range(last_col+"11:"+last_col+"12"," ")
            worksheet.merge_range(last_col+"13:"+last_col+"14"," ")
            worksheet.merge_range(last_col+"15:"+last_col+"16"," ")
            worksheet.merge_range(last_col+"17:"+last_col+"18"," ")
            worksheet.merge_range(last_col+"19:"+last_col+"20"," ")
            worksheet.merge_range(last_col+"21:"+last_col+"23"," ")
            worksheet.merge_range(last_col+"25:"+last_col+"27"," ")
            worksheet.merge_range(last_col+"28:"+last_col+"29"," ")
            worksheet.merge_range(last_col+"30:"+last_col+"31"," ")
            worksheet.merge_range(last_col+"32:"+last_col+"33"," ")
            worksheet.merge_range(last_col+"34:"+last_col+"35"," ")
            worksheet.merge_range(last_col+"36:"+last_col+"37"," ")

            #'''===================='''
            worksheet.merge_range(NoteWholeRow, f"備註：{MONTH}/1日起各隊夜班時間調整，班別顯示於下班日。", last_row_format)
            #'''===================='''  

        table_construction()

        #'''======================================================'''
        #'''*********************** DATA *************************'''


        #'''Data'''
        # 總值勤官, 值勤官
        gen_officer = df_officer.iloc[DAY][0]
        dep_officer = df_officer.iloc[DAY][1]


        tao1_df = tao1_df.replace("9-21","09-21")
        u = tao1_df.iloc

        cell_C9 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D9 = " " if cell_C9 == " " else u[0,:][0][-3:]
        cell_E9 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F9 = " " if cell_E9 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]
        new_time_list = []
        for i in range(len(u[DAY,2:][mask])):
            if "21-" in u[DAY,2:][mask][i]: # 前日21-9
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "0830" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "09-18" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "09-21" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "1130-2130" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        cell_G9 =  leaders_col[new_time_list[0]] # time
        cell_H9 = leaders[new_time_list[0]][-3:]
        cell_G10 = leaders_col[new_time_list[1]]
        cell_H10 = leaders[new_time_list[1]][-3:]



        tao2_df = tao2_df.replace("9-21","09-21")
        u = tao2_df.iloc
        cell_C11 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D11 = " " if cell_C11 == " " else u[0,:][0][-3:]
        cell_E11 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F11 = " " if cell_E11 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]
        for i in range(len(u[DAY,2:][mask])):
            if "21-9" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "0830-1730" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "09-21" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "9-21" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "1130-2130" in u[DAY,2:][mask][i]:
                new_time_list.append(i)

        cell_G11 = leaders_col[new_time_list[0]]#u_[DAY,leaders_col[0]] # time
        cell_H11 = leaders[new_time_list[0]][-3:]
        cell_G12 = leaders_col[new_time_list[1]] #u_[DAY,leaders_col[1]] # time
        cell_H12 = leaders[new_time_list[1]][-3:]


        tao4_df = tao4_df.replace("9-21","09-21")
        u = tao4_df.iloc

        cell_C13 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D13 = " " if cell_C13 == " " else u[0,:][0][-3:]
        cell_E13 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F13 = " " if cell_E13 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        for i in range(len(u[DAY,2:][mask])):
            if "21-9" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "0830-1730" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "09-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "1130-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)

        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]
        cell_G13 = leaders_col[new_time_list[0]]
        cell_H13 = leaders[new_time_list[0]][-3:]
        cell_G14 = leaders_col[new_time_list[1]]
        cell_H14 = leaders[new_time_list[1]][-3:]


        u = tao5_df.iloc

        cell_C15 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D15 = " " if cell_C15 == " " else u[0,:][0][-3:]
        cell_E15 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F15 = " " if cell_E15 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        for i in range(len(u[DAY,2:][mask])):
            if "21-9" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "08-18" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "0830-1730" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "09-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "1130-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)

        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]
        cell_G15 = leaders_col[new_time_list[0]]
        cell_H15 = leaders[new_time_list[0]][-3:]
        cell_G16 = leaders_col[new_time_list[1]]
        cell_H16 = leaders[new_time_list[1]][-3:]


        u = tao3_df.iloc

        cell_C17 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D17 = " " if cell_C17 == " " else u[0,:][0][-3:]
        cell_E17 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F17 = " " if cell_E17 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        for i in range(len(u[DAY,2:][mask])):
            if "08-16" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "11-21" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if u[DAY,2:][mask][i] == "21-09":
                new_time_list.append(i)

        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]
        cell_G17 = leaders_col[new_time_list[0]]
        cell_H17 = leaders[new_time_list[0]][-3:]
        cell_G18 = leaders_col[new_time_list[1]]
        cell_H18 = leaders[new_time_list[1]][-3:]

        cell_G19 = leaders_col[new_time_list[2]]
        cell_H19 = leaders[new_time_list[2]][-3:]
        cell_G20 = leaders_col[new_time_list[3]]
        cell_H20 =  leaders[new_time_list[3]][-3:]



        u = specialForce_df.iloc
        cell_C21 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D21 = " " if cell_C21 == " " else u[0,:][0][-3:]
        cell_E21 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F21 = " " if cell_E21 == " " else u[0,:][1][-2:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]

        # repeat 22-10
        for i in range(len(u[DAY,2:][mask])):
            if "22-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "10-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        if "22-10" in leaders_col.values.tolist():
            new_time_list.append(new_time_list[0])

        cell_G21 = leaders_col[new_time_list[0]]
        cell_H21 = leaders[new_time_list[0]][1:]
        cell_G22 = leaders_col[new_time_list[1]]
        cell_H22 = leaders[new_time_list[1]][1:]



        u = keelung_df.iloc
        cell_C25 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D25 = " " if cell_C25 == " " else keelung_df.columns.tolist()[0][-3:]
        cell_E25 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F25 = " " if cell_E25 == " " else keelung_df.columns.tolist()[1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask].index.tolist()
        # repeat 22-10
        for i in range(len(u[DAY,2:][mask])):
            if "22-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "10-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "12-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        if "22-10" in leaders_col.values.tolist():
            new_time_list.append(new_time_list[0])

        cell_G25 = leaders_col[new_time_list[0]]
        cell_H25 = leaders[new_time_list[0]][-3:]
        cell_G26 = leaders_col[new_time_list[1]]
        cell_H26 = leaders[new_time_list[1]][-3:]



        u = songshang_df.iloc

        cell_C28 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D28 = " " if cell_C28 == " " else " "
        cell_E28 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F28 = " " if cell_E28 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        for i in range(len(u[DAY,2:][mask])):
            if "07" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):     
            if "11-23" in u[DAY,2:][mask][i]:
                new_time_list.append(i)

        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]

        cell_G28 = leaders_col[new_time_list[0]]
        cell_H28 = leaders[new_time_list[0]][-3:]
        cell_G29 = leaders_col[new_time_list[1]]
        cell_H29 = leaders[new_time_list[1]][-3:]



        u = taichung_df.iloc
        cell_C30 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D30 = " " if cell_C30 == " " else u[0,:][0][-3:]
        cell_E30 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F30 = " " if cell_E30 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]
        for i in range(len(u[DAY,2:][mask])):
            if "0745" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):     
            if "09-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)

        cell_G30 =  leaders_col[new_time_list[0]]
        cell_H30 = leaders[new_time_list[0]][-3:]
        cell_G31 = leaders_col[new_time_list[1]]
        cell_H31 = leaders[new_time_list[1]][-3:]


        u = kaohsiungAirport_df.iloc
        cell_C32 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D32 = " " if cell_C32 == " " else u[0,:][0][-3:]
        cell_E32 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F32 = " " if cell_E32 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]
        for i in range(len(u[DAY,2:][mask])):
            if "05" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "14-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)

        cell_G32 = leaders_col[new_time_list[0]]
        cell_H32 = leaders[new_time_list[0]][-3:]
        cell_G33 = leaders_col[new_time_list[1]]
        cell_H33 = leaders[new_time_list[1]][-3:]


        u = kaohsiungPort_df.iloc
        cell_C34 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D34 = " " if cell_C34 == " " else u[0,:][0][-3:]
        cell_E34 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F34 = " " if cell_E34 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]

        for i in range(len(u[DAY,2:][mask])):     
            if "20-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "-16" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):     
            if "-20" in u[DAY,2:][mask][i]:
                new_time_list.append(i)

        if len(new_time_list) != 0 and len(new_time_list) >= 2:
            cell_G34 = leaders_col[new_time_list[0]]
            cell_H34 = leaders[new_time_list[0]][-3:]
            cell_G35 = leaders_col[new_time_list[1]]
            cell_H35 = leaders[new_time_list[1]][-3:]
        
        elif len(new_time_list)!= 0 and len(new_time_list) == 1:
            cell_G34 = leaders_col[new_time_list[0]]
            cell_H34 = leaders[new_time_list[0]][-3:]
            cell_G35 = None
            cell_H35 = None
        elif len(new_time_list) == 0:
            cell_G34 = None
            cell_H34 = None
            cell_G35 = None
            cell_H35 = None


        u = jingmen_df.iloc
        cell_C36 = " " if len(u[DAY,:][0]) <= 2 else u[DAY,:][0].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][0].strip().split("-")[1] # time 
        cell_D36 = " " if cell_C36 == " " else u[0,:][0][-3:]
        cell_E36 = " " if len(u[DAY,:][1]) <= 2 else u[DAY,:][1].strip().split("-")[0]+ "\n|\n"+ u[DAY,:][1].strip().split("-")[1] # time 
        cell_F36 = " " if cell_E36 == " " else u[0,:][1][-3:]

        mask = u[DAY,:] != "輪休"
        new_time_list = []
        leaders_col = u[DAY,2:][mask]
        leaders = u[0,2:][mask]

        for i in range(len(u[DAY,2:][mask])):     
            if "8-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)
        for i in range(len(u[DAY,2:][mask])):
            if "9-" in u[DAY,2:][mask][i]:
                new_time_list.append(i)

        if len(new_time_list) != 0 and len(new_time_list) >= 2:
            cell_G36 = leaders_col[new_time_list[0]]
            cell_H36 = leaders[new_time_list[0]][-3:]
            cell_G37 = leaders_col[new_time_list[1]]
            cell_H37 = leaders[new_time_list[1]][-3:]
        
        elif len(new_time_list)!= 0 and len(new_time_list) == 1:
            cell_G36 = leaders_col[new_time_list[0]]
            cell_H36 = leaders[new_time_list[0]][-3:]
            cell_G37 = None
            cell_H37 = None
        elif len(new_time_list) == 0:
            cell_G36 = None
            cell_H36 = None
            cell_G37 = None
            cell_H37 = None


        #'''********************* WRITE TO EXCEL *************************'''
        #'''Time Data Formatting + Conditional Formatting'''
        # Captains format
        captains_time_format_morning = workbook.add_format({"bold":True,"text_wrap":True,
                                                    "align":"center","valign":"vcenter",
                                                    "bg_color":"#F1C232","font_size":10,
                                                    "border":THICK})

        captains_time_format_not_morning = workbook.add_format({"bold":True,"text_wrap":True,
                                                    "align":"center","valign":"vcenter",
                                                    "bg_color":"#CFAFE7","font_size":10,
                                                    "border":THICK})
                                                    
        text_format_10 = workbook.add_format({"font_size":10,"align":"center",
                                            "valign":"vcenter"})

        text_format = workbook.add_format({"font_size":14,"align":"center",
                                            "valign":"vcenter"})


        # Leaders format
        leaders_time_format_morning = workbook.add_format({"bold":True,"text_wrap":False,
                                                    "align":"center","valign":"vcenter",
                                                    "bg_color":"#F1C232","font_size":10,
                                                    "border":THICK})

        leaders_time_format_not_morning = workbook.add_format({"bold":True,"text_wrap":False,
                                                    "align":"center","valign":"vcenter",
                                                    "bg_color":"#CFAFE7","font_size":10,
                                                    "border":THICK})

        #'''===================== 協管室值勤 ===================='''
        worksheet.merge_range("C4:I4", gen_officer,workbook.add_format({"font_size":16, "align":"center",
                                                "valign":"vcenter","left":2,"right":5,"bottom":2,"top":5}))
        worksheet.merge_range("C5:I5", dep_officer,text_format_center_officers)


        #'''===================== 桃園機場幹部 ===================='''
        TIME_THRESHOLD = 11

        # 桃一隊 tao1_df
        try:
            if int(cell_C9[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C9:C10",cell_C9, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C9:C10",cell_C9, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C9:C10",cell_C9, text_format_10) # time 

        try:
            if int(cell_E9[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E9:E10",cell_E9, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E9:E10",cell_E9, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E9:E10",cell_E9, text_format_10) # time

        # 桃一隊 leaders
        if int(cell_G9[:2]) <= TIME_THRESHOLD:
            worksheet.write("G9", cell_G9, leaders_time_format_morning) # time 
        else:
            if cell_G9[:2] == "21":
                worksheet.write("G9", "昨"+cell_G9, leaders_time_format_not_morning) # time
            else:
                worksheet.write("G9", cell_G9, leaders_time_format_not_morning) # time 

        if int(cell_G10[:2]) <= TIME_THRESHOLD:
            worksheet.write("G10", cell_G10, leaders_time_format_morning) # time 
        else:
            if cell_G10[:2] == "21":
                worksheet.write("G10", "昨"+cell_G10, leaders_time_format_not_morning) # time 
            else:
                worksheet.write("G10", cell_G10, leaders_time_format_not_morning) # time 

        worksheet.merge_range("D9:D10",cell_D9, text_format)  
        worksheet.merge_range("F9:F10",cell_F9, text_format) 
        worksheet.write("H9", cell_H9,text_format)
        worksheet.write("H10", cell_H10, text_format)

        # 桃2隊 tao2_df
        try:
            if int(cell_C11[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C11:C12",cell_C11, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C11:C12",cell_C11, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C11:C12",cell_C11, text_format_10) # time 

        try:
            if int(cell_E11[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E11:E12",cell_E11, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E11:E12",cell_E11, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E11:E12",cell_E11, text_format_10) # time


        if int(cell_G11[:2]) <= TIME_THRESHOLD:
            worksheet.write("G11", cell_G11, leaders_time_format_morning) # time 
        else:
            if cell_G11[:2] == "21":
                worksheet.write("G11", "昨"+cell_G11, leaders_time_format_not_morning) # time 
            else:
                worksheet.write("G11", cell_G11, leaders_time_format_not_morning) # time         

        if int(cell_G12[:2]) <= TIME_THRESHOLD:
            worksheet.write("G12", cell_G12, leaders_time_format_morning) # time 
        else:
            if cell_G12[:2] =="21":
                worksheet.write("G12", "昨"+cell_G12, leaders_time_format_not_morning) # time 
            else:
                worksheet.write("G12", cell_G12, leaders_time_format_not_morning) # time 

        worksheet.merge_range("D11:D12",cell_D11, text_format)
        worksheet.merge_range("F11:F12",cell_F11, text_format)
        worksheet.write("H11", cell_H11,text_format)
        worksheet.write("H12", cell_H12, text_format)

        # 桃4隊 tao4_df
        try:
            if int(cell_C13[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C13:C14",cell_C13, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C13:C14",cell_C13, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C13:C14",cell_C13, text_format_10) # time 

        try:
            if int(cell_E13[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E13:E14",cell_E13, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E13:E14",cell_E13, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E13:E14",cell_E13, text_format_10) # time


        if int(cell_G13[:2]) <= TIME_THRESHOLD:
            worksheet.write("G13", cell_G13, leaders_time_format_morning) # time 
        else:
            if cell_G13[:2] == "21":
                worksheet.write("G13", "昨"+cell_G13, leaders_time_format_not_morning) # time         
            else:
                worksheet.write("G13", cell_G13, leaders_time_format_not_morning) # time 

        if int(cell_G14[:2]) <= TIME_THRESHOLD:
            worksheet.write("G14", cell_G14, leaders_time_format_morning) # time 
        else:
            if cell_G14[:2] == "21":
                worksheet.write("G14", "昨"+cell_G14, leaders_time_format_not_morning) # time 
            else:
                worksheet.write("G14", cell_G14, leaders_time_format_not_morning) # time 

        worksheet.merge_range("D13:D14",cell_D13, text_format)
        worksheet.merge_range("F13:F14",cell_F13, text_format)
        worksheet.write("H13", cell_H13,text_format)
        worksheet.write("H14", cell_H14, text_format)

        # 桃5隊 tao5_df
        try:
            if int(cell_C15[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C15:C16",cell_C15, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C15:C16",cell_C15, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C15:C16",cell_C15, text_format_10) # time 

        try:
            if int(cell_E15[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E15:E16",cell_E15, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E15:E16",cell_E15, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E15:E16",cell_E15, text_format_10) # time

        if int(cell_G15[:2]) <= TIME_THRESHOLD:
            worksheet.write("G15", cell_G15, leaders_time_format_morning) # time 
        else:
            if cell_G15[:2] == "21":
                worksheet.write("G15", "昨"+cell_G15, leaders_time_format_not_morning) # time 
            else:
                worksheet.write("G15", cell_G15, leaders_time_format_not_morning) # time    

        if int(cell_G16[:2]) <= TIME_THRESHOLD:
            worksheet.write("G16", cell_G16, leaders_time_format_morning) # time 
        else:
            if cell_G16[:2] == "21":
                worksheet.write("G16", "昨"+cell_G16, leaders_time_format_not_morning) # time 
            else:
                worksheet.write("G16", cell_G16, leaders_time_format_not_morning) # time 

        worksheet.merge_range("D15:D16",cell_D15, text_format)
        worksheet.merge_range("F15:F16",cell_F15, text_format)
        worksheet.write("H15", cell_H15,text_format)
        worksheet.write("H16", cell_H16, text_format)


        # 桃3隊 tao3_df
        try:
            if int(cell_C17[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C17:C20",cell_C17, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C17:C20",cell_C17, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C17:C20",cell_C17, text_format_10) # time 

        try:
            if int(cell_E17[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E17:E20",cell_E17, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E17:E20",cell_E17, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E17:E20",cell_E17, text_format_10) # time

        if int(cell_G17.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G17", cell_G17, leaders_time_format_morning) # time 
        else:
            worksheet.write("G17", cell_G17, leaders_time_format_not_morning) # time 

        if int(cell_G18.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G18", cell_G18, leaders_time_format_morning) # time 
        else:
            worksheet.write("G18", cell_G18, leaders_time_format_not_morning) # time 

        if int(cell_G19.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G19", cell_G19, leaders_time_format_morning) # time 
        else:
            worksheet.write("G19", cell_G19, leaders_time_format_not_morning) # time 

        if int(cell_G20.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G20", cell_G20, leaders_time_format_morning) # time 
        else:
            if cell_G20[-2:] == "09":
                worksheet.write("G20", "21-明9", leaders_time_format_not_morning) # time 
            else:
                worksheet.write("G20", cell_G20, leaders_time_format_not_morning) # time 


        worksheet.merge_range("D17:D20",cell_D17, text_format)
        worksheet.merge_range("F17:F20",cell_F17, text_format)
        worksheet.write("H17", cell_H17,text_format)
        worksheet.write("H18", cell_H18, text_format)
        worksheet.write("H19", cell_H19, text_format)
        worksheet.write("H20", cell_H20, text_format)

        # 特殊勤務 specialForce_df
        try:
            if int(cell_C21[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C21:C23",cell_C21, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C21:C23",cell_C21, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C21:C23",cell_C21, text_format_10) # time 

        try:
            if int(cell_E21[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E21:E23",cell_E21, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E21:E23",cell_E21, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E21:E23",cell_E21, text_format_10) # time

        if int(cell_G21.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G21", cell_G21, leaders_time_format_morning) # time 
        else:
            if cell_G21[:2] == "22":
                worksheet.write("G21", "昨"+cell_G21, leaders_time_format_not_morning) # time 
            else:
                worksheet.write("G21", cell_G21, leaders_time_format_not_morning) # time 

        if int(cell_G22.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G22", cell_G22, leaders_time_format_morning) # time 
        else:
            worksheet.write("G22", cell_G22, leaders_time_format_not_morning) # time 

        worksheet.write("G23", "22-明10", leaders_time_format_not_morning) # time 
        worksheet.merge_range("D21:D23",cell_D21, text_format)
        worksheet.merge_range("F21:F23",cell_F21, text_format)
        worksheet.write("H21", cell_H21,text_format)
        worksheet.write("H22", cell_H22, text_format)
        worksheet.write("H23", cell_H21, text_format)

        #'''====================== 外機港隊 ======================'''
        # keelunng_df
        try:
            if int(cell_C25[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C25:C27",cell_C25, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C25:C27",cell_C25, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C25:C27",cell_C25, text_format_10) # time 

        try:
            if int(cell_E25[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E25:E27",cell_E25, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E25:E27",cell_E25, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E25:E27",cell_E25, text_format_10) # time
            
        if int(cell_G25.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G25", cell_G25, leaders_time_format_morning) # time 
        else:
            worksheet.write("G25", "昨22-10", leaders_time_format_not_morning) # time 

        if int(cell_G26.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G26", cell_G26, leaders_time_format_morning) # time 
        else:
            worksheet.write("G26", cell_G26, leaders_time_format_not_morning) # time 


        worksheet.write("G27", "22-明10", leaders_time_format_not_morning) # time 


        worksheet.merge_range("D25:D27",cell_D25, text_format)
        worksheet.merge_range("F25:F27",cell_F25, text_format)
        worksheet.write("H25", cell_H25,text_format)
        worksheet.write("H26", cell_H26, text_format)
        worksheet.write("H27", cell_H25, text_format)

        # songshang_df
        try:
            if int(cell_C28[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C28:C29",cell_C28, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C28:C29",cell_C28, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C28:C29",cell_C28, text_format_10) # time 

        try:
            if int(cell_E28[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E28:E29",cell_E28, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E28:E29",cell_E28, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E28:E29",cell_E28, text_format_10) # time

        if int(cell_G28.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G28", cell_G28, leaders_time_format_morning) # time 
        else:
            worksheet.write("G28", cell_G28, leaders_time_format_not_morning) # time 
        if int(cell_G29.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G29", cell_G29, leaders_time_format_morning) # time 
        else:
            worksheet.write("G29", cell_G29, leaders_time_format_not_morning) # time 

        worksheet.merge_range("D28:D29",cell_D28, text_format)
        worksheet.merge_range("F28:F29",cell_F28, text_format)
        worksheet.write("H28", cell_H28,text_format)
        worksheet.write("H29", cell_H29, text_format)

        # taichung_df
        try:
            if int(cell_C30[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C30:C31",cell_C30, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C30:C31",cell_C30, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C30:C31",cell_C30, text_format_10) # time 

        try:
            if int(cell_E30[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E30:E31",cell_E30, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E30:E31",cell_E30, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E30:E31",cell_E30, text_format_10) # time


        if int(cell_G30[:2]) <= TIME_THRESHOLD:
            worksheet.write("G30", cell_G30, leaders_time_format_morning) # time 
        else:
            worksheet.write("G30", cell_G30, leaders_time_format_not_morning) # time 
        if int(cell_G31[:2]) <= TIME_THRESHOLD:
            worksheet.write("G31", cell_G31, leaders_time_format_morning) # time 
        else:
            worksheet.write("G31", cell_G31, leaders_time_format_not_morning) # time 

        worksheet.merge_range("D30:D31",cell_D30, text_format)
        worksheet.merge_range("F30:F31",cell_F30, text_format)
        worksheet.write("H30", cell_H30,text_format)
        worksheet.write("H31", cell_H31, text_format)

        # kaohsiungAirport_df
        try:
            if int(cell_C32[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C32:C33",cell_C32, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C32:C33",cell_C32, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C32:C33",cell_C32, text_format_10) # time 

        try:
            if int(cell_E32[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E32:E33",cell_E32, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E32:E33",cell_E32, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E32:E33",cell_E32, text_format_10) # time

        if int(cell_G32.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G32", cell_G32, leaders_time_format_morning) # time 
        else:
            worksheet.write("G32", cell_G32, leaders_time_format_not_morning) # time 
        if int(cell_G33.split("-")[0]) <= TIME_THRESHOLD:
            worksheet.write("G33", cell_G33, leaders_time_format_morning) # time 
        else:
            worksheet.write("G33", cell_G33, leaders_time_format_not_morning) # time 

        worksheet.merge_range("D32:D33",cell_D32, text_format)
        worksheet.merge_range("F32:F33",cell_F32, text_format)
        worksheet.write("H32", cell_H32,text_format)
        worksheet.write("H33", cell_H33, text_format)

        # kaohsiungPort_df
        try:
            if int(cell_C34[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C34:C35",cell_C34, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C34:C35",cell_C34, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C34:C35",cell_C34, text_format_10) # time 

        try:
            if int(cell_E34[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E34:E35",cell_E34, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E34:E35",cell_E34, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E34:E35",cell_E34, text_format_10) # time


        worksheet.merge_range("D34:D35",cell_D34, text_format)
        worksheet.merge_range("F34:F35",cell_F34, text_format)


        if cell_G34 == None and cell_H34 == None:
            worksheet.merge_range("G34:G35","",grey_bg_color)
            worksheet.merge_range("H34:H35","", grey_bg_color)
        else:
            if cell_G35 == None:
                if int(cell_G34.split("-")[0]) <= TIME_THRESHOLD:
                    worksheet.merge_range("G34:G35", cell_G34, leaders_time_format_morning)
                else:
                    if cell_G34[:2] == "20":
                        worksheet.merge_range("G34:G35", "昨"+cell_G34, leaders_time_format_not_morning)
                    else:
                        worksheet.merge_range("G34:G35", cell_G34, leaders_time_format_not_morning)
                
                worksheet.merge_range("H34:H35", cell_H34,text_format)
            else:
                if int(cell_G34.split("-")[0]) <= TIME_THRESHOLD:
                    worksheet.write("G34", cell_G34, leaders_time_format_morning) # time 
                else:
                    if cell_G34[:2] == "20":
                        worksheet.write("G34:G35", "昨"+cell_G34, leaders_time_format_not_morning)
                    else:
                        worksheet.write("G34:G35", cell_G34, leaders_time_format_not_morning)

                if int(cell_G35.split("-")[0]) <= TIME_THRESHOLD:
                    worksheet.write("G35", cell_G35, leaders_time_format_morning) # time 
                else:
                    if cell_G35[:2] == "20":
                        worksheet.write("G35", "昨"+cell_G35, leaders_time_format_not_morning) # time 
                    else:
                        worksheet.write("G35", cell_G35, leaders_time_format_not_morning) # time 

                worksheet.write("H34", cell_H34,text_format)
                worksheet.write("H35", cell_H35, text_format)
        
        # jingmen_df
        try:
            if int(cell_C36[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("C36:C37",cell_C36, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("C36:C37",cell_C36, captains_time_format_not_morning)

        except:
            worksheet.merge_range("C36:C37",cell_C36, text_format_10) # time 

        try:
            if int(cell_E36[:2]) <= TIME_THRESHOLD:
                worksheet.merge_range("E36:E37",cell_E36, captains_time_format_morning) # time 
            else:
                worksheet.merge_range("E36:E37",cell_E36, captains_time_format_not_morning)
        except:
            worksheet.merge_range("E36:E37",cell_E36, text_format_10) # time

        worksheet.merge_range("D36:D37",cell_D36, text_format)
        worksheet.merge_range("F36:F37",cell_F36, text_format)

        if cell_G36 == None and cell_H36 == None:
            worksheet.merge_range("G36:G37","",)
            worksheet.merge_range("H36:H37","", )
        else:
            if cell_G37 == None:
                if int(cell_G36.split("-")[0]) <= TIME_THRESHOLD:
                    worksheet.merge_range("G36:G37", cell_G36, leaders_time_format_morning)
                else:
                    worksheet.merge_range("G36:G37", cell_G36, leaders_time_format_not_morning)
                
                if cell_H36 == "長王愷":
                    worksheet.merge_range("H36:H37", "王愷",text_format)
                else:
                    worksheet.merge_range("H36:H37", cell_H36,text_format)
            else:
                if int(cell_G36.split("-")[0]) <= TIME_THRESHOLD:
                    worksheet.write("G36", cell_G36, leaders_time_format_morning) # time 
                else:
                    worksheet.write("G36", cell_G36, leaders_time_format_not_morning) # time 
                if int(cell_G37.split("-")[0]) <= TIME_THRESHOLD:
                    worksheet.write("G37", cell_G37, leaders_time_format_morning) # time 
                else:
                    worksheet.write("G37", cell_G37, leaders_time_format_not_morning) # time 

                if cell_H36 == "長王愷":
                    worksheet.write("H36", "王愷",text_format)
                else:
                    worksheet.write("H36", cell_H36,text_format)

                if cell_H37 == "長王愷":
                    worksheet.write("H37", "王愷",text_format)
                else:
                    worksheet.write("H37", cell_H37, text_format)


        worksheet.conditional_format('C11:C12',{'type':"text",'criteria':"not containing",
                                            "value":"|","format":grey_bg_color,
                                            "multi_range":"C11:C12 E11:E12 C15:C16 E15:E16 C25:C27 E25:E27 C30:C31 E30:E31 C34:C35 E34:E35"})

        # End of writing to excel
    
    
    st.markdown("### 3. Download Output")

    # Add a placeholder
    latest_iteration = st.empty()
    bar = st.progress(0)

    for i in range(100):
        # Update the progress bar with each iteration.
        latest_iteration.text(f'Processing ... {i+1} %')
        bar.progress(i + 1)
        time.sleep(0.001)

    message = st.success('Done!')
    if message:

        st.download_button(
            label=f"Download_幹部出勤表{MONTH}月{DAY}日.xlsx",
            data=output.getvalue(),
            file_name=excel_file_name,
            mime="application/vnd.ms-excel"
            )



